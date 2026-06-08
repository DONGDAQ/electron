#!/usr/bin/env python3
"""幻塔月度结算单生成脚本"""

import sys
import os
import shutil
from datetime import datetime, timedelta
from pathlib import Path
from copy import copy, deepcopy

import openpyxl
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from quote_system.feishu_client import FeishuClient, SPREADSHEET_TOKEN
from quote_system.paths import get_settlement_dir

# 单价对照表
UNIT_PRICES = {
    '翻译+润色': 0.64,
    '翻译+角色': 0.64,
    '润色': 0.5,
}

BASE_DIR = Path(__file__).resolve().parent.parent.parent
TEMPLATE_DIR = BASE_DIR / "app" / "模板" / "结算模板"


def excel_serial_to_date(serial):
    """Excel日期序列号 -> datetime对象"""
    base = datetime(1899, 12, 30)
    return base + timedelta(days=float(serial))


def format_date_ymd(dt):
    """日期 -> Y.M.D格式（不补零）如 2026.4.1"""
    return f"{dt.year}.{dt.month}.{dt.day}"


def get_unit_price(content):
    """根据需求内容匹配单价"""
    if not content:
        return 0.64
    for key, price in UNIT_PRICES.items():
        if key in content:
            return price
    return 0.64


def read_feishu_data(year, month):
    """从飞书读取指定月份的交付数据"""
    client = FeishuClient('ZXBogz')
    data = client.read_sheet('ZXBogz', 'A1:O')

    # 计算月份serial范围
    month_start = datetime(year, month, 1)
    if month == 12:
        month_end = datetime(year + 1, 1, 1)
    else:
        month_end = datetime(year, month + 1, 1)
    base = datetime(1899, 12, 30)
    serial_start = (month_start - base).days
    serial_end = (month_end - base).days

    records = []
    for i, row in enumerate(data):
        if i == 0:
            continue
        g_val = row[6]  # G列: 交付日期
        if g_val is None:
            continue
        try:
            g_num = float(g_val)
            if not (serial_start <= g_num < serial_end):
                continue
        except (TypeError, ValueError):
            continue

        records.append({
            'req_name': str(row[3] or ''),      # D列: 文本
            'task_no': str(row[4] or ''),       # E列: Task No
            'content': str(row[2] or ''),       # C列: 需求内容
            'deliv_date': g_num,                # G列: 交付日期
            'order_date': row[5],               # F列: 委托日期
            'row_index': i + 1,
            'source_chars': row[14],            # O列: 源语言字数
            'billable_words': row[13],          # N列: 计费字数
            'amount': row[8],                   # I列: 报价金额
        })

    # 按交付日期排序
    records.sort(key=lambda r: r['deliv_date'])
    return records


def generate_settlement_excel(records, year, month, output_dir):
    """生成结算单Excel文件"""
    template = TEMPLATE_DIR / '【幻塔结算单】-TOF【游戏内】年月模板.xlsx'
    filename = f'【幻塔结算单】-TOF【游戏内】{year}年{month}月.xlsx'
    output_path = output_dir / filename

    # 如果目标文件已存在，先删除（避免文件锁定问题）
    if output_path.exists():
        try:
            os.remove(str(output_path))
        except PermissionError:
            ts = datetime.now().strftime('%H%M%S')
            filename = f'【幻塔结算单】-TOF【游戏内】{year}年{month}月_{ts}.xlsx'
            output_path = output_dir / filename

    shutil.copy2(str(template), str(output_path))
    wb = openpyxl.load_workbook(str(output_path))
    ws = wb.active

    # 设置C4日期
    today = datetime.now()
    ws['C4'] = f'日期：{today.strftime("%Y.%m.%d")}'

    num_records = len(records)
    data_start_row = 11

    total_row = None
    for row in range(1, ws.max_row + 1):
        for col in range(1, 16):
            val = ws.cell(row=row, column=col).value
            if val and '费用合计' in str(val):
                total_row = row
                break
        if total_row:
            break
    if total_row is None:
        total_row = 33

    template_slots = total_row - data_start_row

    # 获取模板行格式（用第12行作为空白模板行的参考）
    ref_row_data = {}
    for col in range(1, 16):
        cell = ws.cell(row=12, column=col)
        ref_row_data[col] = {
            'font': copy(cell.font),
            'fill': copy(cell.fill),
            'border': copy(cell.border),
            'alignment': copy(cell.alignment),
            'number_format': cell.number_format,
        }

    # 如果记录数超过模板行数，插入新行
    if num_records > template_slots:
        extra = num_records - template_slots
        insert_pos = total_row  # 在合计行之前插入
        for _ in range(extra):
            ws.insert_rows(insert_pos)
            # 复制格式到新行
            for col in range(1, 16):
                cell = ws.cell(row=insert_pos, column=col)
                fmt = ref_row_data.get(col, {})
                if fmt:
                    cell.font = copy(fmt.get('font'))
                    cell.fill = copy(fmt.get('fill'))
                    cell.border = copy(fmt.get('border'))
                    cell.alignment = copy(fmt.get('alignment'))
                    cell.number_format = fmt.get('number_format', '')
            insert_pos += 1

    # 填充数据
    for i, rec in enumerate(records):
        row = data_start_row + i
        deliv_dt = excel_serial_to_date(rec['deliv_date'])

        # C列: 委托文本 (Feishu D)
        ws.cell(row=row, column=3, value=rec['req_name'])
        # D列: Task No (Feishu E)
        ws.cell(row=row, column=4, value=rec['task_no'])
        # E列: 需求内容描述 (Feishu C)
        ws.cell(row=row, column=5, value=rec['content'])
        # F列: 交付日期 (Feishu G)
        ws.cell(row=row, column=6, value=deliv_dt)
        # G列: 源语言字数 (Feishu O)
        if rec['source_chars'] is not None:
            ws.cell(row=row, column=7, value=float(rec['source_chars']))
        # H列: 计费字数 (Feishu N)
        if rec['billable_words'] is not None:
            ws.cell(row=row, column=8, value=float(rec['billable_words']))
        # I列: 单价
        ws.cell(row=row, column=9, value=get_unit_price(rec['content']))
        # J列: 费用小计 (Feishu I)
        if rec['amount'] is not None:
            ws.cell(row=row, column=10, value=float(rec['amount']))

    last_data_row = data_start_row + num_records - 1
    for row in range(last_data_row + 1, total_row):
        ws.row_dimensions[row].hidden = True

    wb.save(str(output_path))
    print(f'[OK] 结算单已保存: {output_path}')
    return output_path


def generate_acceptance_docx(records, year, month, output_dir):
    """生成验收单Word文件"""
    template = TEMPLATE_DIR / '【幻塔】验收单-TOF【游戏内】年月模板.docx'
    filename = f'【幻塔】验收单-TOF【游戏内】{year}年{month}月.docx'
    output_path = output_dir / filename

    # 如果目标文件已存在，先删除（避免文件锁定问题）
    if output_path.exists():
        try:
            os.remove(str(output_path))
        except PermissionError:
            from datetime import datetime as dt
            ts = dt.now().strftime('%H%M%S')
            filename = f'【幻塔】验收单-TOF【游戏内】{year}年{month}月_{ts}.docx'
            output_path = output_dir / filename

    shutil.copy2(str(template), str(output_path))
    doc = Document(str(output_path))

    # 更新Date和Number
    today = datetime.now()
    date_str = f"{today.year}/{today.month}/{today.day}"
    # Date段落是P[3]
    p_date = doc.paragraphs[3]
    p_date.clear()
    run_label = p_date.add_run('Date')
    run_label.bold = True
    run_label.font.size = Pt(10.5)
    run_colon = p_date.add_run(':')
    run_colon.font.size = Pt(10.5)
    run_val = p_date.add_run(date_str)
    run_val.font.size = Pt(10.5)

    # Number段落是P[4] - 格式yyyymmdd02
    number_str = today.strftime('%Y%m%d') + '01'
    p_num = doc.paragraphs[4]
    p_num.clear()
    run_label2 = p_num.add_run('Number')
    run_label2.bold = True
    run_label2.font.size = Pt(10.5)
    run_colon2 = p_num.add_run(':')
    run_colon2.font.size = Pt(10.5)
    run_val2 = p_num.add_run(number_str)
    run_val2.font.size = Pt(10.5)

    # 处理表格
    table = doc.tables[0]
    tbl = table._tbl

    # 调整列宽: No.列缩小, Source Language列加大
    # 目标宽度(pct): col0=300, col3=800
    col_widths_pct = {0: 300, 3: 800}
    for ri in range(len(table.rows)):
        row = table.rows[ri]
        for ci, w_val in col_widths_pct.items():
            if ci < len(row.cells):
                tcPr = row.cells[ci]._tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is not None:
                        tcW.set(qn('w:w'), str(w_val))

    # 保存模板数据行(row 3)作为克隆参考, 然后删除多余数据行
    template_tr = table.rows[3]._tr
    # 保留行0-3(标题+表头+模板行), 删除行4-24(剩余数据行)
    while len(table.rows) > 5:  # row0+row1+row2+row3+TOTAL = 5行
        tr_elem = table.rows[4]._tr
        tbl.remove(tr_elem)
    # 现在: row0=title, row1=header, row2=sub-header, row3=模板行, row4=TOTAL

    # 填充第一行数据(row 3)
    _fill_acceptance_row(table.rows[3], records[0], 1)

    # 为剩余记录克隆模板行
    total_tr = table.rows[4]._tr
    for i in range(1, len(records)):
        new_tr = deepcopy(template_tr)
        tbl.insert(list(tbl).index(total_tr), new_tr)
        # 新行现在是倒数第二行(在TOTAL之前)
        new_row_idx = len(table.rows) - 2
        _fill_acceptance_row(table.rows[new_row_idx], records[i], i + 1)

    # 更新TOTAL行的Cost
    total_amount = sum(float(r['amount'] or 0) for r in records)
    total_row = table.rows[-1]
    cost_cell = total_row.cells[9]
    for p in cost_cell.paragraphs:
        p.clear()
        run = p.add_run(str(round(total_amount, 2)))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

    doc.save(str(output_path))
    print(f'[OK] 验收单已保存: {output_path}')
    return output_path


def _fill_acceptance_row(row, rec, row_num):
    """填充验收单表格的一行数据, 只修改文本内容"""
    values = {
        0: str(row_num),
        1: 'TOF',
        2: 'Translation',
        3: 'Chinese (Simplified)',
        4: 'Japanese',
        5: rec['task_no'],
        6: (format_date_ymd(excel_serial_to_date(rec['order_date']))
            if rec['order_date'] is not None else ''),
        7: format_date_ymd(excel_serial_to_date(rec['deliv_date'])),
        8: str(rec['billable_words']) if rec['billable_words'] is not None else '',
        9: str(rec['amount']) if rec['amount'] is not None else '',
        10: 'CNY',
    }
    for ci, text in values.items():
        cell = row.cells[ci]
        # 找到第一个 w:t 元素设置文本
        t_elem = cell._tc.find('.//' + qn('w:t'))
        if t_elem is not None:
            t_elem.text = text
            # 清除其他 w:t 元素中的文本
            for other_t in cell._tc.findall('.//' + qn('w:t')):
                if other_t != t_elem:
                    other_t.text = ''


def main():
    import argparse
    parser = argparse.ArgumentParser(description='幻塔月度结算单生成')
    parser.add_argument('--year', type=int, default=2026, help='年份')
    parser.add_argument('--month', type=int, default=4, help='月份')
    parser.add_argument('--output', type=str, default=None, help='输出目录')
    args = parser.parse_args()

    year, month = args.year, args.month
    output_dir = Path(args.output) if args.output else get_settlement_dir() / f"{year}年{month}月" / "完美世界" / "幻塔"
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f'读取飞书数据: {year}年{month}月...')
    records = read_feishu_data(year, month)
    print(f'找到 {len(records)} 条交付记录')

    if not records:
        print('没有找到匹配的交付记录，退出。')
        return

    generate_settlement_excel(records, year, month, output_dir)
    generate_acceptance_docx(records, year, month, output_dir)
    print('完成！')


if __name__ == '__main__':
    main()
