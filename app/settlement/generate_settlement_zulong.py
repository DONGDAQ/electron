"""祖龙客服月度结算 — 以闪亮之名 / 龙族"""

import re
import sys
import os
from datetime import datetime
from pathlib import Path

from openpyxl import load_workbook

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from quote_system.paths import get_settlement_dir

BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_DIR = BASE_DIR / "模板" / "结算模板"

PROJECTS = {
    "yishan": {
        "name": "以闪亮之名",
        "xlsx_template": "【结算单】日语客服-以闪亮之名-2026年4月.xlsx",
        "docx_template": "国内执行单模板-以闪亮之名-2026年4月.docx",
    },
    "longzu": {
        "name": "龙族",
        "xlsx_template": "【结算单】日语客服-龙族-2026年4月.xlsx",
        "docx_template": "国内执行单模板-龙族-2026年4月.docx",
    },
}


def generate_settlement_xlsx(project_key: str, count: int, year: int, month: int, output_dir: Path) -> Path:
    info = PROJECTS[project_key]
    template_path = TEMPLATE_DIR / info["xlsx_template"]

    wb = load_workbook(template_path)
    ws = wb.active

    today_str = datetime.now().strftime("%Y%m%d")
    ws["D9"] = int(today_str)
    original_c16 = str(ws["C16"].value or "")
    ws["C16"] = re.sub(r"\d+月", f"{month}月", original_c16)
    ws["D16"] = count

    output_name = f"【结算单】日语客服-{info['name']}-{year}年{month}月.xlsx"
    output_path = output_dir / output_name
    wb.save(output_path)
    wb.close()
    return output_path


def generate_acceptance_docx(project_key: str, count: int, year: int, month: int, output_dir: Path) -> Path:
    from docx import Document

    info = PROJECTS[project_key]
    template_path = TEMPLATE_DIR / info["docx_template"]

    doc = Document(str(template_path))

    table = doc.tables[0]
    exec_time_cell = table.cell(2, 1)
    exec_time_para = exec_time_cell.paragraphs[0]
    exec_time_text = exec_time_para.text
    exec_time_para.text = exec_time_text.replace(
        exec_time_text, f"{year}年{month}月"
    )

    exec_content_cell = table.cell(3, 1)
    exec_content_para = exec_content_cell.paragraphs[0]
    exec_content_para.text = f"日语客服-{info['name']}-{month}月份费用"

    output_name = f"国内执行单模板-{info['name']}-{year}年{month}月.docx"
    output_path = output_dir / output_name
    doc.save(str(output_path))
    return output_path


def generate_zulong_settlement(year: int, month: int, counts: dict[str, int]) -> list[dict]:
    month_str = f"{year}年{month}月"
    output_dir = get_settlement_dir() / month_str / "祖龙"
    output_dir.mkdir(parents=True, exist_ok=True)

    files = []
    for key in ["yishan", "longzu"]:
        count = counts.get(key, 0)
        if count <= 0:
            continue
        xlsx_path = generate_settlement_xlsx(key, count, year, month, output_dir)
        docx_path = generate_acceptance_docx(key, count, year, month, output_dir)
        files.append({"name": xlsx_path.name, "path": str(xlsx_path)})
        files.append({"name": docx_path.name, "path": str(docx_path)})

    return files
