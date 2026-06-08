"""Bilibili 项目盖章版结算单生成：对账单确认后，生成盖章用的结算单 PDF。
支持马娘(umamusume)、邦邦2(bang2)、炽焰天穹(hbr)。
"""
from __future__ import annotations

import argparse
import calendar
import os
import shutil
import sys
from datetime import datetime
from pathlib import Path

import openpyxl
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from quote_system.paths import get_settlement_dir

TEMPLATE_DIR = ROOT / "模板" / "结算模板"
SETTLEMENT_COMPANY = "Bilibili"
TEMPLATE_NAME = "bilibili结算单盖章用模板.docx"

SEALED_CONFIG = {
    "umamusume": {"project": "马娘", "project_full_name": "优俊少女", "code": "代号PD"},
    "bang2": {"project": "bang2", "project_full_name": "邦邦2", "code": "邦邦2"},
    "hbr": {"project": "HBR", "project_full_name": "炽焰天穹", "code": "HBR"},
}


def _get_sealed_config(project_key: str) -> dict:
    cfg = SEALED_CONFIG.get(project_key)
    if not cfg:
        raise ValueError(f"不支持的盖章结算项目: {project_key}")
    return cfg


def _recalc_and_read(bill_path: Path) -> float | None:
    """用 Excel 打开文件重算公式，保存后读取 GRAND TOTAL。"""
    import pythoncom
    pythoncom.CoInitialize()
    import win32com.client
    excel = win32com.client.DispatchEx("Excel.Application")
    excel.Visible = False
    excel.DisplayAlerts = False
    wb = None
    try:
        wb = excel.Workbooks.Open(str(bill_path.resolve()))
        wb.Save()
        wb.Close()
        excel.Quit()

        wb2 = openpyxl.load_workbook(bill_path, data_only=True)
        ws = wb2.active
        for r in range(1, ws.max_row + 1):
            v = ws.cell(row=r, column=8).value
            if v and "GRAND TOTAL" in str(v):
                val = ws.cell(row=r, column=9).value
                if isinstance(val, (int, float)) and val > 0:
                    return float(val)
                break
    except Exception:
        pass
    finally:
        try:
            if wb:
                wb.Close()
        except Exception:
            pass
        try:
            excel.Quit()
        except Exception:
            pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass
    return None


def _read_total_from_bill(output_dir: Path, year: int, month: int, code: str) -> float:
    """从对账单 xlsx 读取总金额。优先读公式缓存，没有再调 Excel 重算。"""
    candidates = sorted(
        list(output_dir.glob(f"大连游者之家翻译账单-哔哩哔哩游戏【{code}】-{year}年{month}月*.xlsx")),
        key=lambda p: p.stat().st_mtime, reverse=True
    )
    if not candidates:
        raise FileNotFoundError(f"找不到对账单: {output_dir}")

    for bill_path in candidates:
        # 先尝试直接读 GRAND TOTAL 缓存值
        wb = openpyxl.load_workbook(bill_path, data_only=True)
        ws = wb.active
        for r in range(1, ws.max_row + 1):
            v = ws.cell(row=r, column=8).value
            if v and "GRAND TOTAL" in str(v):
                val = ws.cell(row=r, column=9).value
                if isinstance(val, (int, float)) and val > 0:
                    return float(val)
                break

        # 缓存为空，调 Excel 重算公式后再读
        val = _recalc_and_read(bill_path)
        if val is not None:
            return val

    raise ValueError("无法从对账单中读取总金额")


def generate_sealed(
    year: int,
    month: int,
    total_amount: float | None = None,
    *,
    project_key: str = "umamusume",
    output_dir: Path | None = None,
) -> Path:
    cfg = _get_sealed_config(project_key)
    code = cfg["code"]
    project_name = cfg["project"]
    project_full_name = cfg["project_full_name"]

    template = TEMPLATE_DIR / TEMPLATE_NAME
    if not template.exists():
        raise FileNotFoundError(f"找不到模板: {template}")

    if output_dir is None:
        output_dir = get_settlement_dir() / f"{year}年{month}月" / SETTLEMENT_COMPANY / project_name
    output_dir.mkdir(parents=True, exist_ok=True)

    if total_amount is None:
        total_amount = _read_total_from_bill(output_dir, year, month, code)

    docx_path = _unique_path(output_dir, f"大连游者之家翻译结算单-哔哩哔哩游戏【{code}】-{year}年{month}月.docx")
    pdf_path = docx_path.with_suffix(".pdf")

    shutil.copy2(str(template), str(docx_path))
    doc = Document(str(docx_path))
    table = doc.tables[0]

    _, last_day = calendar.monthrange(year, month)
    amount_str = f"{total_amount:.2f}"

    _set_cell_text(table.cell(5, 2), f"{year}年{month}月1日-{month}月{last_day}日")
    _set_cell_text(table.cell(6, 2), f"{project_full_name}{month}月本地化翻译费用 {amount_str}")
    _set_cell_text(table.cell(8, 2), f"{amount_str} 元")
    today = datetime.now()
    _update_date_runs(table.cell(14, 0), today.year, today.month, today.day)

    doc.save(str(docx_path))

    _convert_to_pdf(docx_path, pdf_path)

    docx_path.unlink()

    return pdf_path


def _unique_path(dir: Path, filename: str) -> Path:
    p = dir / filename
    if not p.exists():
        return p
    stem = p.stem
    for idx in range(1, 100):
        alt = p.with_name(f"{stem}_{idx}{p.suffix}")
        if not alt.exists():
            return alt
    return p


def _set_cell_text(cell, text: str) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    p = cell.paragraphs[0]
    p.add_run(text)


def _update_date_runs(cell, year: int, month: int, day: int) -> None:
    for p in cell.paragraphs:
        if "日期：" not in (p.text or ""):
            continue
        runs = p.runs
        digit_indices = [i for i, r in enumerate(runs) if r.text.strip().isdigit()]
        if len(digit_indices) >= 3:
            runs[digit_indices[0]].text = str(year)
            runs[digit_indices[1]].text = str(month)
            runs[digit_indices[2]].text = str(day)
        return


def _convert_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    """使用 Word COM 对象将 docx 另存为 PDF。"""
    import pythoncom
    pythoncom.CoInitialize()
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False

    doc = None
    try:
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)  # 17 = wdFormatPDF
        doc.Close()
    finally:
        if doc:
            try:
                doc.Close()
            except Exception:
                pass
        word.Quit()
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def main():
    parser = argparse.ArgumentParser(description="Bilibili 项目盖章版结算单生成")
    parser.add_argument("--year", type=int, required=True, help="结算年份")
    parser.add_argument("--month", type=int, required=True, help="结算月份")
    parser.add_argument("--project", type=str, default="umamusume", help="项目: umamusume/bang2/hbr")
    parser.add_argument("--amount", type=float, default=None, help="结算总金额（不传自动从对账单读取）")
    args = parser.parse_args()

    output = generate_sealed(args.year, args.month, args.amount, project_key=args.project)
    print(f"盖章版结算单已生成: {output}")
    os.startfile(str(output))


if __name__ == "__main__":
    main()
