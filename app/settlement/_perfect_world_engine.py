#!/usr/bin/env python3
"""完美世界项目结算单统一引擎 — 幻塔/异环游戏内/异环发行共用"""

import os
import shutil
from copy import copy, deepcopy
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from pathlib import Path

import openpyxl
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from quote_system.feishu_client import FeishuClient
from quote_system.paths import get_settlement_dir


@dataclass
class SettlementProfile:
    """结算单配置"""
    sheet_id: str
    unit_prices: dict[str, float]
    excel_template: str          # 模板文件名（相对于 TEMPLATE_DIR）
    excel_filename: str          # 输出文件名模板 (用 {year}, {month})
    excel_conflict_prefix: str   # 冲突时的备用文件名前缀
    docx_template: str
    docx_filename: str
    docx_conflict_prefix: str
    output_subdir: str           # 输出子目录路径 (相对于 get_settlement_dir()/{year}年{month}月/)
    project_code: str            # 验收单表格 col 1 项目代号
    number_suffix: str           # 验收单 Number 后缀 (01/02/03)
    has_task_no: bool = True     # 飞书是否有 Task No 列
    description: str = "月度结算单生成"
    default_price: float = 0.64


PERFECT_WORLD_PROFILES: dict[str, SettlementProfile] = {
    "huanta": SettlementProfile(
        sheet_id="ZXBogz",
        unit_prices={'翻译+润色': 0.64, '翻译+角色': 0.64, '润色': 0.5},
        excel_template='【幻塔结算单】-TOF【游戏内】年月模板.xlsx',
        excel_filename='【幻塔结算单】-TOF【游戏内】{year}年{month}月.xlsx',
        excel_conflict_prefix='【幻塔结算单】-TOF【游戏内】',
        docx_template='【幻塔】验收单-TOF【游戏内】年月模板.docx',
        docx_filename='【幻塔】验收单-TOF【游戏内】{year}年{month}月.docx',
        docx_conflict_prefix='【幻塔】验收单-TOF【游戏内】',
        output_subdir='完美世界/幻塔',
        project_code='TOF',
        number_suffix='01',
        description='幻塔月度结算单生成',
    ),
    "yihuan_nei": SettlementProfile(
        sheet_id="bfba7c",
        unit_prices={'翻译+润色': 0.64, '翻译+角色': 0.64, '润色': 0.5},
        excel_template='【异环结算单】-NTE【游戏内】年月模板.xlsx',
        excel_filename='【异环结算单】-NTE【游戏内】{year}年{month}月.xlsx',
        excel_conflict_prefix='【异环游戏内结算单】-NTE【游戏内】',
        docx_template='【异环】验收单-NTE【游戏内】年月模板.docx',
        docx_filename='【异环】验收单-NTE【游戏内】{year}年{month}月.docx',
        docx_conflict_prefix='【异环游戏内】验收单-TOF【游戏内】',
        output_subdir='完美世界/异环游戏内',
        project_code='NTE',
        number_suffix='02',
        description='异环游戏内月度结算单生成',
    ),
    "yihuan_faxing": SettlementProfile(
        sheet_id="S5yHmP",
        unit_prices={'翻译+润色英日': 0.72, '翻译+润色': 0.64, '翻译+角色': 0.64, '润色': 0.5},
        excel_template='【异环结算单】-NTE【发行】模板.xlsx',
        excel_filename='【异环结算单】-NTE【发行】{year}年{month}月.xlsx',
        excel_conflict_prefix='【异环发行结算单】-NTE【发行】',
        docx_template='【异环】验收单-NTE【发行】模板.docx',
        docx_filename='【异环】验收单-NTE【发行】{year}年{month}月.docx',
        docx_conflict_prefix='【异环发行】验收单-TOF【发行】',
        output_subdir='完美世界/异环发行',
        project_code='NTE',
        number_suffix='03',
        has_task_no=False,
        description='异环发行月度结算单生成',
    ),
}

BASE_DIR = Path(__file__).resolve().parent.parent.parent
TEMPLATE_DIR = BASE_DIR / "app" / "模板" / "结算模板"


def _excel_serial_to_date(serial):
    base = datetime(1899, 12, 30)
    return base + timedelta(days=float(serial))


def _format_date_ymd(dt):
    return f"{dt.year}.{dt.month}.{dt.day}"


def _get_unit_price(content, prices, default_price):
    if not content:
        return default_price
    for key, price in prices.items():
        if key in content:
            return price
    return default_price


def read_feishu_data(year, month, profile: SettlementProfile):
    """从飞书读取指定月份的交付数据"""
    client = FeishuClient(profile.sheet_id)
    data = client.read_sheet(profile.sheet_id, 'A1:O')

    month_start = datetime(year, month, 1)
    if month == 12:
        month_end = datetime(year + 1, 1, 1)
    else:
        month_end = datetime(year, month + 1, 1)
    base = datetime(1899, 12, 30)
    serial_start = (month_start - base).days
    serial_end = (month_end - base).days

    records = []
    for i, row in enumerate(data):
        if i == 0:
            continue
        g_val = row[6]
        if g_val is None:
            continue
        try:
            g_num = float(g_val)
            if not (serial_start <= g_num < serial_end):
                continue
        except (TypeError, ValueError):
            continue

        records.append({
            'req_name': str(row[3] or ''),
            'task_no': str(row[4] or '') if profile.has_task_no else '',
            'content': str(row[2] or ''),
            'deliv_date': g_num,
            'order_date': row[5],
            'row_index': i + 1,
            'source_chars': row[14],
            'billable_words': row[13],
            'amount': row[8],
        })

    records.sort(key=lambda r: r['deliv_date'])
    return records


def _get_safe_path(output_dir: Path, filename: str, conflict_prefix: str):
    """获取安全的输出路径，避免文件锁定冲突"""
    output_path = output_dir / filename
    if output_path.exists():
        try:
            os.remove(str(output_path))
        except PermissionError:
            ts = datetime.now().strftime('%H%M%S')
            output_path = output_dir / f"{conflict_prefix}_{ts}.xlsx" if output_path.suffix == '.xlsx' else output_dir / f"{conflict_prefix}_{ts}.docx"
    return output_path


def generate_settlement_excel(records, year, month, output_dir, profile: SettlementProfile):
    """生成结算单Excel文件"""
    template = TEMPLATE_DIR / profile.excel_template
    filename = profile.excel_filename.format(year=year, month=month)
    output_path = _get_safe_path(output_dir, filename, profile.excel_conflict_prefix)

    shutil.copy2(str(template), str(output_path))
    wb = openpyxl.load_workbook(str(output_path))
    ws = wb.active

    today = datetime.now()
    ws['C4'] = f'日期：{today.strftime("%Y.%m.%d")}'

    num_records = len(records)
    data_start_row = 11

    total_row = None
    for row in range(1, ws.max_row + 1):
        for col in range(1, 16):
            val = ws.cell(row=row, column=col).value
            if val and '费用合计' in str(val):
                total_row = row
                break
        if total_row:
            break
    if total_row is None:
        total_row = 33

    template_slots = total_row - data_start_row

    ref_row_data = {}
    for col in range(1, 16):
        cell = ws.cell(row=12, column=col)
        ref_row_data[col] = {
            'font': copy(cell.font),
            'fill': copy(cell.fill),
            'border': copy(cell.border),
            'alignment': copy(cell.alignment),
            'number_format': cell.number_format,
        }

    if num_records > template_slots:
        extra = num_records - template_slots
        insert_pos = total_row
        for _ in range(extra):
            ws.insert_rows(insert_pos)
            for col in range(1, 16):
                cell = ws.cell(row=insert_pos, column=col)
                fmt = ref_row_data.get(col, {})
                if fmt:
                    cell.font = copy(fmt.get('font'))
                    cell.fill = copy(fmt.get('fill'))
                    cell.border = copy(fmt.get('border'))
                    cell.alignment = copy(fmt.get('alignment'))
                    cell.number_format = fmt.get('number_format', '')
            insert_pos += 1

    prices = profile.unit_prices
    default_price = profile.default_price
    for i, rec in enumerate(records):
        row = data_start_row + i
        deliv_dt = _excel_serial_to_date(rec['deliv_date'])

        ws.cell(row=row, column=3, value=rec['req_name'])
        ws.cell(row=row, column=4, value=rec['task_no'] if profile.has_task_no else '-')
        ws.cell(row=row, column=5, value=rec['content'])
        ws.cell(row=row, column=6, value=deliv_dt)
        if rec['source_chars'] is not None:
            ws.cell(row=row, column=7, value=float(rec['source_chars']))
        if rec['billable_words'] is not None:
            ws.cell(row=row, column=8, value=float(rec['billable_words']))
        ws.cell(row=row, column=9, value=_get_unit_price(rec['content'], prices, default_price))
        if rec['amount'] is not None:
            ws.cell(row=row, column=10, value=float(rec['amount']))

    last_data_row = data_start_row + num_records - 1
    for row in range(last_data_row + 1, total_row):
        ws.row_dimensions[row].hidden = True

    wb.save(str(output_path))
    print(f'[OK] 结算单已保存: {output_path}')
    return output_path


def _fill_acceptance_row(row, rec, row_num, profile: SettlementProfile):
    """填充验收单表格的一行数据"""
    values = {
        0: str(row_num),
        1: profile.project_code,
        2: 'Translation',
        3: 'Chinese (Simplified)',
        4: 'Japanese',
        5: rec['task_no'] if profile.has_task_no else rec['req_name'],
        6: (_format_date_ymd(_excel_serial_to_date(rec['order_date']))
            if rec['order_date'] is not None else ''),
        7: _format_date_ymd(_excel_serial_to_date(rec['deliv_date'])),
        8: str(rec['billable_words']) if rec['billable_words'] is not None else '',
        9: str(rec['amount']) if rec['amount'] is not None else '',
        10: 'CNY',
    }
    for ci, text in values.items():
        cell = row.cells[ci]
        t_elem = cell._tc.find('.//' + qn('w:t'))
        if t_elem is not None:
            t_elem.text = text
            for other_t in cell._tc.findall('.//' + qn('w:t')):
                if other_t != t_elem:
                    other_t.text = ''


def generate_acceptance_docx(records, year, month, output_dir, profile: SettlementProfile):
    """生成验收单Word文件"""
    template = TEMPLATE_DIR / profile.docx_template
    filename = profile.docx_filename.format(year=year, month=month)
    output_path = _get_safe_path(output_dir, filename, profile.docx_conflict_prefix)

    shutil.copy2(str(template), str(output_path))
    doc = Document(str(output_path))

    today = datetime.now()
    date_str = f"{today.year}/{today.month}/{today.day}"
    p_date = doc.paragraphs[3]
    p_date.clear()
    run_label = p_date.add_run('Date')
    run_label.bold = True
    run_label.font.size = Pt(10.5)
    run_colon = p_date.add_run(':')
    run_colon.font.size = Pt(10.5)
    run_val = p_date.add_run(date_str)
    run_val.font.size = Pt(10.5)

    number_str = today.strftime('%Y%m%d') + profile.number_suffix
    p_num = doc.paragraphs[4]
    p_num.clear()
    run_label2 = p_num.add_run('Number')
    run_label2.bold = True
    run_label2.font.size = Pt(10.5)
    run_colon2 = p_num.add_run(':')
    run_colon2.font.size = Pt(10.5)
    run_val2 = p_num.add_run(number_str)
    run_val2.font.size = Pt(10.5)

    table = doc.tables[0]
    tbl = table._tbl

    col_widths_pct = {0: 300, 3: 800}
    for ri in range(len(table.rows)):
        row = table.rows[ri]
        for ci, w_val in col_widths_pct.items():
            if ci < len(row.cells):
                tcPr = row.cells[ci]._tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is not None:
                        tcW.set(qn('w:w'), str(w_val))

    template_tr = table.rows[3]._tr
    while len(table.rows) > 5:
        tr_elem = table.rows[4]._tr
        tbl.remove(tr_elem)

    _fill_acceptance_row(table.rows[3], records[0], 1, profile)

    total_tr = table.rows[4]._tr
    for i in range(1, len(records)):
        new_tr = deepcopy(template_tr)
        tbl.insert(list(tbl).index(total_tr), new_tr)
        new_row_idx = len(table.rows) - 2
        _fill_acceptance_row(table.rows[new_row_idx], records[i], i + 1, profile)

    total_amount = sum(float(r['amount'] or 0) for r in records)
    total_row = table.rows[-1]
    cost_cell = total_row.cells[9]
    for p in cost_cell.paragraphs:
        p.clear()
        run = p.add_run(str(round(total_amount, 2)))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

    doc.save(str(output_path))
    print(f'[OK] 验收单已保存: {output_path}')
    return output_path


def run_settlement(profile_key: str, year: int, month: int, output_dir: Path | None = None):
    """统一入口：运行指定项目的月度结算"""
    profile = PERFECT_WORLD_PROFILES[profile_key]
    if output_dir is None:
        output_dir = get_settlement_dir() / f"{year}年{month}月" / profile.output_subdir
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f'[{profile.description}] 读取飞书数据: {year}年{month}月...')
    records = read_feishu_data(year, month, profile)
    print(f'找到 {len(records)} 条交付记录')

    if not records:
        print('没有找到匹配的交付记录，退出。')
        return None, None

    excel_path = generate_settlement_excel(records, year, month, output_dir, profile)
    docx_path = generate_acceptance_docx(records, year, month, output_dir, profile)
    print('完成！')
    return excel_path, docx_path
